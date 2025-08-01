#!/usr/bin/env python3
"""
Gerador de documentos DOC - Substituto do gerador PDF
Autor: Claude Code
Data: 01/08/2025
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor
import os
from datetime import datetime
import json

class GeradorDocumentos:
    """Gerador de documentos DOC para relatórios NPS"""
    
    def __init__(self):
        # Garante que os relatórios sejam salvos na pasta correta (projeto-Z/relatorios)
        self.base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # Volta para pasta raiz do projeto
        self.relatorios_dir = os.path.abspath(os.path.join(self.base_dir, "relatorios"))
        print(f"[DEBUG] Pasta de relatórios configurada para salvar em: {self.relatorios_dir}")
        print(f"[DEBUG] Diretório base do projeto: {self.base_dir}")
        self._criar_pasta_relatorios()
    
    def _criar_pasta_relatorios(self):
        """Cria pasta de relatórios se não existir"""
        try:
            if not os.path.exists(self.relatorios_dir):
                os.makedirs(self.relatorios_dir)
                print(f"[DEBUG] Pasta de relatórios criada em: {self.relatorios_dir}")
            else:
                print(f"[DEBUG] Pasta de relatórios já existe em: {self.relatorios_dir}")
            if not os.path.exists(self.relatorios_dir):
                os.makedirs(self.relatorios_dir)
                print(f"[DEBUG] Pasta de relatórios criada: {self.relatorios_dir}")
            else:
                print(f"[DEBUG] Pasta de relatórios já existe: {self.relatorios_dir}")
        except Exception as e:
            print(f"[ERROR] Erro ao criar pasta de relatórios: {e}")
    
    def _configurar_estilos(self, doc):
        """Configura estilos personalizados para o documento"""
        try:
            # Estilo para título principal
            titulo_style = doc.styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
            titulo_font = titulo_style.font
            titulo_font.name = 'Calibri'
            titulo_font.size = Pt(18)
            titulo_font.bold = True
            titulo_font.color.rgb = RGBColor(31, 78, 121)  # Azul corporativo
            titulo_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            titulo_style.paragraph_format.space_after = Pt(12)
            
            # Estilo para subtítulos
            subtitulo_style = doc.styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
            subtitulo_font = subtitulo_style.font
            subtitulo_font.name = 'Calibri'
            subtitulo_font.size = Pt(14)
            subtitulo_font.bold = True
            subtitulo_font.color.rgb = RGBColor(68, 114, 196)
            subtitulo_style.paragraph_format.space_before = Pt(12)
            subtitulo_style.paragraph_format.space_after = Pt(6)
            
            # Estilo para métricas destacadas
            metrica_style = doc.styles.add_style('Metrica', WD_STYLE_TYPE.PARAGRAPH)
            metrica_font = metrica_style.font
            metrica_font.name = 'Calibri'
            metrica_font.size = Pt(12)
            metrica_font.bold = True
            metrica_font.color.rgb = RGBColor(0, 176, 80)  # Verde
            metrica_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            print(f"[DOC] Aviso: Erro ao criar estilos personalizados: {e}")
    
    def gerar_relatorio_nps(self, dados_analise, nome_loja, nome_arquivo=None):
        """Gera relatório completo de NPS em formato DOC"""
        try:
            # Criar documento
            doc = Document()
            self._configurar_estilos(doc)
            
            # Nome do arquivo
            if not nome_arquivo:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                nome_arquivo = f"relatorio_nps_{nome_loja.replace(' ', '_')}_{timestamp}.docx"
            
            # Garantir extensão .docx
            if not nome_arquivo.endswith('.docx'):
                nome_arquivo += '.docx'
            
            caminho_completo = os.path.join(self.relatorios_dir, nome_arquivo)
            print(f"[DEBUG] Salvando documento em: {caminho_completo}")
            
            # CABEÇALHO DO RELATÓRIO
            titulo = doc.add_paragraph(f"RELATÓRIO DE ANÁLISE NPS", style='TituloPrincipal')
            
            # Informações básicas
            doc.add_paragraph(f"Loja: {nome_loja}", style='Subtitulo')
            doc.add_paragraph(f"Data de Geração: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
            doc.add_paragraph("")  # Linha em branco
            
            # RESUMO EXECUTIVO
            doc.add_paragraph("RESUMO EXECUTIVO", style='Subtitulo')
            
            # Se dados_analise for string (análise IA), processar como texto
            if isinstance(dados_analise, str):
                self._processar_analise_texto(doc, dados_analise)
            
            # Se for dicionário, processar estrutura
            elif isinstance(dados_analise, dict):
                self._processar_analise_estruturada(doc, dados_analise)
            
            # Se for lista, processar como dados tabulares
            elif isinstance(dados_analise, list):
                self._processar_dados_tabulares(doc, dados_analise)
            
            else:
                # Fallback: converter para string
                doc.add_paragraph(str(dados_analise))
            
            # RODAPÉ
            doc.add_paragraph("")
            doc.add_paragraph("─" * 60)
            rodape = doc.add_paragraph(f"Relatório gerado automaticamente pelo DashBot")
            rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Salvar documento
            doc.save(caminho_completo)
            print(f"[DOC] Relatório gerado: {nome_arquivo}")
            print(f"[DEBUG] Caminho completo do arquivo: {caminho_completo}")
            print(f"[DEBUG] Arquivo existe após salvar: {os.path.exists(caminho_completo)}")
            print(f"[DEBUG] Conteúdo da pasta relatorios:")
            try:
                for arquivo in os.listdir(self.relatorios_dir):
                    print(f"  - {arquivo}")
            except Exception as e:
                print(f"[ERROR] Erro ao listar pasta: {e}")
            
            return caminho_completo
            
        except Exception as e:
            print(f"[ERROR] Erro ao gerar documento DOC: {e}")
            return None
    
    def _processar_analise_texto(self, doc, texto_analise):
        """Processa análise em formato de texto (IA)"""
        try:
            # Dividir texto em seções
            secoes = texto_analise.split('\n\n')
            
            for secao in secoes:
                if not secao.strip():
                    continue
                
                linhas = secao.strip().split('\n')
                primeira_linha = linhas[0].strip()
                
                # Se parece com título (termina com :, está em maiúscula, etc.)
                if (primeira_linha.endswith(':') or 
                    primeira_linha.isupper() or 
                    any(palavra in primeira_linha.upper() for palavra in 
                        ['ANÁLISE', 'RESUMO', 'CONCLUSÃO', 'RECOMENDAÇÃO', 'MÉTRICA'])):
                    
                    # Adicionar como subtítulo
                    doc.add_paragraph(primeira_linha, style='Subtitulo')
                    
                    # Adicionar resto como conteúdo
                    if len(linhas) > 1:
                        conteudo = '\n'.join(linhas[1:])
                        doc.add_paragraph(conteudo)
                else:
                    # Adicionar como parágrafo normal
                    doc.add_paragraph(secao.strip())
                
                doc.add_paragraph("")  # Espaço entre seções
                
        except Exception as e:
            print(f"[DOC] Erro ao processar texto: {e}")
            doc.add_paragraph(str(texto_analise))
    
    def _processar_analise_estruturada(self, doc, dados_dict):
        """Processa análise em formato de dicionário"""
        try:
            # Métricas principais
            if 'nps_score' in dados_dict:
                doc.add_paragraph(f"NPS Score: {dados_dict['nps_score']:.1f}%", style='Metrica')
            
            if 'total_registros' in dados_dict:
                doc.add_paragraph(f"Total de Respostas: {dados_dict['total_registros']}")
            
            if 'avg_rating' in dados_dict:
                doc.add_paragraph(f"Avaliação Média: {dados_dict['avg_rating']:.1f}/10")
            
            doc.add_paragraph("")
            
            # Processar outras chaves
            for chave, valor in dados_dict.items():
                if chave not in ['nps_score', 'total_registros', 'avg_rating']:
                    if isinstance(valor, (dict, list)):
                        doc.add_paragraph(f"{chave.replace('_', ' ').title()}:", style='Subtitulo')
                        doc.add_paragraph(str(valor))
                    else:
                        doc.add_paragraph(f"{chave.replace('_', ' ').title()}: {valor}")
            
        except Exception as e:
            print(f"[DOC] Erro ao processar estrutura: {e}")
            doc.add_paragraph(str(dados_dict))
    
    def _processar_dados_tabulares(self, doc, dados_lista):
        """Processa dados em formato de lista/tabela"""
        try:
            doc.add_paragraph("DADOS ANALISADOS", style='Subtitulo')
            doc.add_paragraph(f"Total de registros: {len(dados_lista)}")
            doc.add_paragraph("")
            
            # Se a lista contém dicionários, criar uma tabela simples
            if dados_lista and isinstance(dados_lista[0], dict):
                # Pegar primeiros 5 registros como exemplo
                amostra = dados_lista[:5]
                
                for i, registro in enumerate(amostra, 1):
                    doc.add_paragraph(f"Registro {i}:")
                    for chave, valor in registro.items():
                        doc.add_paragraph(f"  • {chave}: {valor}")
                    doc.add_paragraph("")
                
                if len(dados_lista) > 5:
                    doc.add_paragraph(f"... e mais {len(dados_lista) - 5} registros")
            
            else:
                # Lista simples
                for item in dados_lista[:10]:  # Primeiros 10 itens
                    doc.add_paragraph(f"• {item}")
                
                if len(dados_lista) > 10:
                    doc.add_paragraph(f"... e mais {len(dados_lista) - 10} itens")
                    
        except Exception as e:
            print(f"[DOC] Erro ao processar dados tabulares: {e}")
            doc.add_paragraph(f"Dados processados: {len(dados_lista)} registros")

# Função de compatibilidade com o sistema existente
def gerar_doc_inteligente(dados_analise, nome_loja, nome_arquivo=None):
    """Função compatível com o sistema existente para gerar DOC"""
    try:
        gerador = GeradorDocumentos()
        return gerador.gerar_relatorio_nps(dados_analise, nome_loja, nome_arquivo)
    except Exception as e:
        print(f"[ERROR] Erro na função gerar_doc_inteligente: {e}")
        return None

# Função de teste
def main():
    """Função de teste do gerador de documentos"""
    try:
        # Dados de teste
        dados_teste = {
            'nps_score': 75.5,
            'total_registros': 150,
            'avg_rating': 8.2,
            'analise': 'Resultado satisfatório com boa aceitação dos clientes.'
        }
        
        resultado = gerar_doc_inteligente(dados_teste, "Loja Teste")
        
        if resultado:
            print(f"[TEST] Documento de teste gerado: {resultado}")
        else:
            print("[TEST] Erro ao gerar documento de teste")
            
    except Exception as e:
        print(f"[ERROR] Erro no teste: {e}")

if __name__ == "__main__":
    main()